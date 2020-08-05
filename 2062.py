from docx import Document
from copy import deepcopy
from openpyxl import Workbook
from openpyxl import load_workbook

def import_data(wb_name):
    #Open excel table, read in each specific value to the data dictionary, each item becomes a list of 12 values, one for each column
    wb=load_workbook(wb_name)
    ws=wb.active
    #items=[['a','b','c','d','e','f','A','B','C','D','E','F'] for i in range(100)]
    #data={'from':'0','to':'1','hand_receipt_number':'2','end_item_stock_number':'3','end_item_description':'4','publication_number':'5','publication_date':'6','quantity':'7','items':items}
    max_row = ws.max_row
    print(("Copying {} items from property.xlsx to da2062_out.docx").format(max_row))
    data['from']=str(ws.cell(1,2).value)
    data['to']=str(ws.cell(2,2).value)
    data['hand_receipt_number']=str(ws.cell(3,2).value)
    data['end_item_stock_number']=str(ws.cell(4,2).value)
    data['end_item_description']=str(ws.cell(5,2).value)
    data['publication_number']=str(ws.cell(6,2).value)
    data['publication_date']=str(ws.cell(7,2).value)
    data['quantity']=str(ws.cell(8,2).value)
    items=[]
    for i in range(10, max_row + 1):
        item_row=[str(ws.cell(i,c).value) for c in range(1,13)]
        items.append(item_row)
    data['items']=items
    #uncomment to override with test data
    #items=[['a','b','c','d','e','f','A','B','C','D','E','F'] for i in range(100)]
    #data={'from':'0','to':'1','hand_receipt_number':'2','end_item_stock_number':'3','end_item_description':'4','publication_number':'5','publication_date':'6','quantity':'7','items':items}
    return data

def replace_text(cell, par_num,run_num,new_text,old_text=""):
    #replaces a string in a single run while maintaining the original run's formatting
    p=cell.paragraphs[par_num]
    inline = p.runs[run_num]
    text = inline.text.replace(old_text, new_text,1)
    inline.text = text

def add_text(cell,new_text):
    #concatenates text to the end of a run. Ignores "None" because this is the value read in from blank excel cells
    if new_text!="None":
        p=cell.paragraphs[-1]
        inline = p.runs[4]
        inline.text =new_text

def add_pages(doc,num_new_pages):
    #Copies the table on the second page and adds it to the end of the file as many times as needed.
    subtext=doc.paragraphs[-1]
    newsubtext = deepcopy(subtext)
    for page in range(num_new_pages):
        template = doc.tables[1]
        tbl = template._tbl
        # Here we do the copy of the table
        new_tbl = deepcopy(tbl)
        # Then we do the replacement
        paragraph = doc.add_paragraph()
        # After that, we add the previously copied table
        paragraph._p.addnext(new_tbl)
        paragraph = doc.add_paragraph()
        renum_page(newsubtext,paragraph,page,num_new_pages)

def renum_page(newsubtext,paragraph,page,num_new_pages):
    #Modifies the page numbers at the end of new pages while keeping the formatting.
    i=0
    for run in newsubtext.runs:
        if i==10:
            output_run = paragraph.add_run(str(page+3),run.style)
        elif i==21:
            output_run = paragraph.add_run(str(num_new_pages+2),run.style)
        else:
            output_run = paragraph.add_run(run.text,run.style)
        if i==0:
            for _ in range(11):
                output_run.add_tab()
        elif i==2:
            output_run.add_tab()
        # Run's bold data
        output_run.bold = run.bold
        # Run's italic data
        output_run.italic = run.italic
        # Run's underline data
        output_run.underline = run.underline
        # Run's font data
        output_run.font.name=run.font.name
        output_run.font.size=run.font.size
        i+=1


#there 16 rows on the first page and 19 on the second page

def add_items(items,doc):
    #Adds each item row to the document
    '''a_col=0
    b_col=2
    c_col=7
    d_col=8
    e_col=9
    f_col=10
    g_A_col=11
    g_B_col=12
    g_C_col=13
    g_D_col=14
    g_E_col=15
    g_F_col=16'''
    first_table_cols=[0,2,7,8,9,10,11,12,13,14,15,16] #The first page has these columns
    subs_table_cols=[i for i in range(12)] #Every other page has 0-11 columns
    table_cols=first_table_cols
    table_row=4
    tab_num=0
    for i in range(len(items)):
        item=items[i]
        if (i+3)%19==0:
            table_cols=subs_table_cols
            tab_num+=1
            table_row=2
        tab=doc.tables[tab_num]
        item_col=0
        for table_col in table_cols:
            add_text(tab.cell(table_row,table_col),item[item_col])
            item_col+=1
        table_row+=1

def fill_form(data):
    #Fills in available admin data, calculates the required number of pages, changes the page count on the first page
    doc = Document('da2062.docx')
    tab=doc.tables[0] #a list of all tables in document
    from_cell=tab.cell(0,4)
    add_text(from_cell,data["from"])

    to_cell=tab.cell(0,6)
    add_text(to_cell,data["to"])

    hand_receipt_number_cell=tab.cell(0,14)
    add_text(hand_receipt_number_cell,data["hand_receipt_number"])

    end_item_stock_number_cell=tab.cell(1,1)
    add_text(end_item_stock_number_cell,data["end_item_stock_number"])

    end_item_description_cell=tab.cell(1,3)
    add_text(end_item_description_cell,data["end_item_description"])

    publication_number_cell=tab.cell(1,5)
    add_text(publication_number_cell,data["publication_number"])

    publication_date_cell=tab.cell(1,10)
    add_text(publication_date_cell,data["publication_date"])

    quantity_cell=tab.cell(1,14)
    add_text(quantity_cell,data["quantity"])
    
    req_pages=(len(data['items'])+2)//19+1

    replace_text(doc,-1,10,"2")
    replace_text(doc,-1,21,str(req_pages))

    if req_pages>2:
        add_pages(doc,req_pages-2)
    add_items(data['items'],doc)

    page_num_cell=tab.cell(20,0)
    replace_text(page_num_cell,-1,-3,str(req_pages),"1")
    return doc

def main():
    data=import_data("property.xlsx")
    doc=fill_form(data)
    doc.save('da2062_out.docx')

main()
